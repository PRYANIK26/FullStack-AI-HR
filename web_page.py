import streamlit as st
import os
from dotenv import load_dotenv
from openai import OpenAI
import json
import docx
import PyPDF2
from striprtf.striprtf import rtf_to_text
import plotly.graph_objects as go
import plotly.express as px
import sqlite3
import uuid
from datetime import datetime
import pandas as pd
import re
import requests


MAIN_SERVER_URL = f"http://localhost:{os.getenv('MAIN_SERVER_PORT', 5000)}/approve-interview"

load_dotenv()

st.set_page_config(
    page_title="HR-Аватар: Анализ кандидатов",
    page_icon="🤖",
    layout="wide"
)

if 'current_view' not in st.session_state:
    st.session_state.current_view = 'main'
if 'selected_analysis_id' not in st.session_state:
    st.session_state.selected_analysis_id = None
if 'analysis_mode' not in st.session_state:
    st.session_state.analysis_mode = 'single'

class DatabaseManager:
    """Управление базой данных с возможностью переключения SQLite/PostgreSQL"""
    
    def __init__(self, db_type='sqlite'):
        self.db_type = db_type
        if db_type == 'sqlite':
            self.db_path = 'hr_analysis.db'
            self._init_sqlite()
        elif db_type == 'postgresql':
            self._init_postgresql()
    
    def _init_sqlite(self):
        """Инициализация SQLite"""
        self.conn = sqlite3.connect(self.db_path, check_same_thread=False)
        self.conn.row_factory = sqlite3.Row
        self._create_tables()
    
    def _init_postgresql(self):
        """Инициализация PostgreSQL (для Railway)"""
        import psycopg2
        from psycopg2.extras import RealDictCursor
        
        DATABASE_URL = os.getenv('DATABASE_URL')
        self.conn = psycopg2.connect(DATABASE_URL, cursor_factory=RealDictCursor)
        self._create_tables()
    
    def _create_tables(self):
        """Создание таблиц базы данных"""
        cursor = self.conn.cursor()
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS vacancies (
                id TEXT PRIMARY KEY,
                title TEXT NOT NULL,
                content TEXT NOT NULL,
                industry TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS candidates (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                resume_content TEXT NOT NULL,
                email TEXT,
                phone TEXT,
                telegram TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS analyses (
                id TEXT PRIMARY KEY,
                vacancy_id TEXT NOT NULL,
                candidate_id TEXT NOT NULL,
                result_json TEXT NOT NULL,
                total_score INTEGER,
                recommendation TEXT,
                confidence_level TEXT,
                analysis_type TEXT DEFAULT 'single',
                needs_interview BOOLEAN DEFAULT FALSE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (vacancy_id) REFERENCES vacancies (id),
                FOREIGN KEY (candidate_id) REFERENCES candidates (id)
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS interviews (
                id TEXT PRIMARY KEY,
                analysis_id TEXT NOT NULL,
                candidate_name TEXT NOT NULL,
                vacancy_title TEXT NOT NULL,
                start_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                end_time TIMESTAMP,
                duration_seconds INTEGER,
                total_questions INTEGER DEFAULT 0,
                interview_plan TEXT,  -- JSON
                final_scores TEXT,    -- JSON
                final_recommendation TEXT,
                interviewer_notes TEXT,
                phase_breakdown TEXT, -- JSON
                adaptive_insights TEXT, -- JSON
                repetition_analysis TEXT, -- JSON - НОВОЕ ПОЛЕ
                timing_statistics TEXT, -- JSON - НОВОЕ ПОЛЕ  
                advanced_analytics TEXT, -- JSON - НОВОЕ ПОЛЕ
                status TEXT DEFAULT 'active',
                FOREIGN KEY (analysis_id) REFERENCES analyses (id)
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS interview_qa (
                id TEXT PRIMARY KEY,
                interview_id TEXT NOT NULL,
                question_number INTEGER,
                question_area TEXT,
                question_text TEXT NOT NULL,
                question_difficulty TEXT,
                question_phase TEXT,
                answer_text TEXT,
                answer_duration_seconds INTEGER,
                technical_score INTEGER,
                communication_score INTEGER,
                depth_score INTEGER,
                confidence_score INTEGER,
                practical_experience INTEGER,
                red_flags TEXT,  -- JSON
                strengths_shown TEXT,  -- JSON
                analysis_notes TEXT,
                knowledge_gaps TEXT,  -- JSON - НОВОЕ ПОЛЕ
                adaptation_needed TEXT,  -- НОВОЕ ПОЛЕ
                repetition_detected BOOLEAN DEFAULT FALSE,  -- НОВОЕ ПОЛЕ
                alternative_strategy_used TEXT,  -- НОВОЕ ПОЛЕ
                timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (interview_id) REFERENCES interviews (id)
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS interview_timing (
                id TEXT PRIMARY KEY,
                interview_id TEXT NOT NULL,
                question_number INTEGER,
                question_start_time TIMESTAMP,
                answer_duration_seconds INTEGER,
                analysis_duration_seconds INTEGER,
                phase TEXT,
                time_status TEXT,  -- on_track, need_acceleration, critical_time
                remaining_minutes INTEGER,
                FOREIGN KEY (interview_id) REFERENCES interviews (id)
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS final_evaluations (
                id TEXT PRIMARY KEY,
                analysis_id TEXT NOT NULL UNIQUE,
                evaluation_summary TEXT,
                final_recommendation TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (analysis_id) REFERENCES analyses (id)
            )
        ''')

        try: cursor.execute('ALTER TABLE candidates ADD COLUMN email TEXT')
        except: pass
        try: cursor.execute('ALTER TABLE candidates ADD COLUMN phone TEXT')
        except: pass
        try: cursor.execute('ALTER TABLE candidates ADD COLUMN telegram TEXT')
        except: pass
        try: cursor.execute('ALTER TABLE analyses ADD COLUMN needs_interview BOOLEAN DEFAULT FALSE')
        except: pass
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS batch_analyses (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                analysis_type TEXT NOT NULL,
                total_analyses INTEGER DEFAULT 0,
                completed_analyses INTEGER DEFAULT 0,
                results_json TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        self.conn.commit()
    
    def save_vacancy(self, title, content, industry=None):
        """Сохранение вакансии"""
        vacancy_id = str(uuid.uuid4())
        cursor = self.conn.cursor()
        cursor.execute(
            "INSERT INTO vacancies (id, title, content, industry) VALUES (?, ?, ?, ?)",
            (vacancy_id, title, content, industry)
        )
        self.conn.commit()
        return vacancy_id
    
    def save_candidate(self, name, resume_content, email=None, phone=None, telegram=None):
        """Сохранение кандидата с контактными данными"""
        candidate_id = str(uuid.uuid4())
        cursor = self.conn.cursor()
        cursor.execute(
            "INSERT INTO candidates (id, name, resume_content, email, phone, telegram) VALUES (?, ?, ?, ?, ?, ?)",
            (candidate_id, name, resume_content, email, phone, telegram)
        )
        self.conn.commit()
        return candidate_id
    
    def update_candidate_contacts(self, candidate_id, email=None, phone=None, telegram=None):
        """Обновление контактных данных кандидата"""
        cursor = self.conn.cursor()
        cursor.execute(
            "UPDATE candidates SET email = ?, phone = ?, telegram = ? WHERE id = ?",
            (email, phone, telegram, candidate_id)
        )
        self.conn.commit()
    
    def save_analysis(self, vacancy_id, candidate_id, result_json, total_score, recommendation, confidence_level, analysis_type='single', needs_interview=False):
        """Сохранение результата анализа"""
        analysis_id = str(uuid.uuid4())
        cursor = self.conn.cursor()
        cursor.execute(
            "INSERT INTO analyses (id, vacancy_id, candidate_id, result_json, total_score, recommendation, confidence_level, analysis_type, needs_interview) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (analysis_id, vacancy_id, candidate_id, json.dumps(result_json), total_score, recommendation, confidence_level, analysis_type, needs_interview)
        )
        self.conn.commit()
        return analysis_id
    
    def update_interview_status(self, analysis_id, needs_interview):
        """Обновление статуса собеседования"""
        cursor = self.conn.cursor()
        cursor.execute(
            "UPDATE analyses SET needs_interview = ? WHERE id = ?",
            (needs_interview, analysis_id)
        )
        self.conn.commit()
    
    def save_batch_analysis(self, name, analysis_type, results):
        """Сохранение batch анализа"""
        batch_id = str(uuid.uuid4())
        cursor = self.conn.cursor()
        cursor.execute(
            "INSERT INTO batch_analyses (id, name, analysis_type, total_analyses, completed_analyses, results_json) VALUES (?, ?, ?, ?, ?, ?)",
            (batch_id, name, analysis_type, len(results), len(results), json.dumps(results))
        )
        self.conn.commit()
        return batch_id
    
    def get_all_analyses(self):
        """Получение всех анализов с информацией о наличии интервью"""
        cursor = self.conn.cursor()
        cursor.execute('''
            SELECT 
                a.id, 
                a.total_score,
                a.recommendation,
                a.created_at,
                v.title as vacancy_title, 
                c.name as candidate_name,
                MAX(i.id) as interview_id, -- Используем MAX чтобы сгруппировать дубликаты интервью
                fe.id as final_evaluation_id
            FROM analyses a
            JOIN vacancies v ON a.vacancy_id = v.id
            JOIN candidates c ON a.candidate_id = c.id
            LEFT JOIN interviews i ON a.id = i.analysis_id AND i.status = 'completed'
            LEFT JOIN final_evaluations fe ON a.id = fe.analysis_id
            GROUP BY a.id -- ГРУППИРУЕМ по ID анализа, чтобы убрать дубликаты
            ORDER BY a.created_at DESC
        ''')
        return cursor.fetchall()
    
    def get_analysis_by_id(self, analysis_id):
        """Получение анализа по ID"""
        cursor = self.conn.cursor()
        cursor.execute('''
            SELECT a.*, v.title as vacancy_title, v.content as vacancy_content,
                   c.name as candidate_name, c.resume_content as candidate_resume,
                   c.email as candidate_email, c.phone as candidate_phone, c.telegram as candidate_telegram
            FROM analyses a
            JOIN vacancies v ON a.vacancy_id = v.id
            JOIN candidates c ON a.candidate_id = c.id
            WHERE a.id = ?
        ''', (analysis_id,))
        return cursor.fetchone()

    def get_interview_by_analysis_id(self, analysis_id):
        """Получение отчета по интервью по ID анализа"""
        cursor = self.conn.cursor()
        cursor.execute('''
            SELECT * FROM interviews WHERE analysis_id = ? AND status = 'completed'
        ''', (analysis_id,))
        interview_data = cursor.fetchone()

        if not interview_data:
            return None, None
        
        cursor.execute('''
            SELECT * FROM interview_qa WHERE interview_id = ? ORDER BY question_number ASC
        ''', (interview_data['id'],))
        qa_data = cursor.fetchall()
        
        return dict(interview_data), [dict(row) for row in qa_data]

    def search_analyses(self, search_term):
        """Поиск анализов с информацией о наличии интервью"""
        cursor = self.conn.cursor()
        cursor.execute('''
            SELECT 
                a.*, 
                v.title as vacancy_title, 
                c.name as candidate_name,
                i.id as interview_id,
                fe.id as final_evaluation_id
            FROM analyses a
            JOIN vacancies v ON a.vacancy_id = v.id
            JOIN candidates c ON a.candidate_id = c.id
            LEFT JOIN interviews i ON a.id = i.analysis_id AND i.status = 'completed'
            LEFT JOIN final_evaluations fe ON a.id = fe.analysis_id
            WHERE v.title LIKE ? OR c.name LIKE ? OR a.recommendation LIKE ?
            ORDER BY a.created_at DESC
        ''', (f'%{search_term}%', f'%{search_term}%', f'%{search_term}%'))
        return cursor.fetchall()

    def save_final_evaluation(self, analysis_id, summary, recommendation):
        """Сохранение финального заключения"""
        eval_id = str(uuid.uuid4())
        cursor = self.conn.cursor()
        cursor.execute(
            "INSERT OR REPLACE INTO final_evaluations (id, analysis_id, evaluation_summary, final_recommendation) VALUES (?, ?, ?, ?)",
            (eval_id, analysis_id, summary, recommendation)
        )
        self.conn.commit()

    def get_final_evaluation(self, analysis_id):
        """Получение финального заключения"""
        cursor = self.conn.cursor()
        cursor.execute("SELECT * FROM final_evaluations WHERE analysis_id = ?", (analysis_id,))
        return cursor.fetchone()

class CandidateEvaluator:
    def __init__(self):
        api_key = os.getenv('OPENAI_API_KEY')
        if not api_key:
            st.error("OPENAI_API_KEY не найден в .env файле!")
            st.stop()
            
        try:
            self.client = OpenAI(api_key=api_key)
        except TypeError:
            import openai
            openai.api_key = api_key
            self.client = None
    
    def classify_industry(self, vacancy_text):
        """Определение отрасли через отдельный API вызов"""
        classification_prompt = f"""
Определи отрасль для следующей вакансии из строго этого списка:
- tech (IT, разработка, программирование, DevOps, архитектура ПО)
- fintech (финансовые технологии, банки, платежи, финтех, криптовалюты)
- healthcare (медицина, здравоохранение, фармацевтика, биотехнологии)
- retail (розничная торговля, e-commerce, продажи, маркетинг товаров)
- manufacturing (производство, промышленность, заводы, логистика)
- consulting (консалтинг, стратегия, бизнес-анализ, управленческие услуги)
- education (образование, обучение, академия, университеты, школы)
- government (государственные услуги, муниципальные органы, госструктуры)

ВАЖНО: 
- Учитывай ОСНОВНУЮ специализацию позиции, а не второстепенные упоминания
- Если в IT-вакансии упоминается "обучение команды" - это все равно tech, а не education
- Если в финтех-вакансии упоминается "разработка" - это fintech, а не tech

Текст вакансии:
{vacancy_text}

Ответь ТОЛЬКО название отрасли одним словом:"""

        try:
            response = self.client.chat.completions.create(
                model="gpt-5-mini",
                messages=[
                    {"role": "user", "content": classification_prompt}
                ]
            )
            industry = response.choices[0].message.content.strip().lower()
            
            valid_industries = ['tech', 'fintech', 'healthcare', 'retail', 'manufacturing', 
                              'consulting', 'education', 'government']
            
            if industry in valid_industries:
                return industry
            else:
                return 'tech'
                
        except Exception as e:
            st.warning(f"Ошибка классификации отрасли: {e}. Используется общий анализ.")
            return 'tech'
    
    def extract_text_from_file(self, uploaded_file):
        """Извлечение текста из файлов"""
        try:
            uploaded_file.seek(0)
            file_name = uploaded_file.name.lower()
            
            if file_name.endswith('.docx'):
                doc = docx.Document(uploaded_file)
                text_parts = []
                
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        text_parts.append(text)
                
                for table in doc.tables:
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_data.append(cell_text)
                        if row_data:
                            text_parts.append(" | ".join(row_data))
                
                result = "\n".join(text_parts)
                return result if result else "Файл пустой"
                
            elif file_name.endswith('.pdf'):
                uploaded_file.seek(0)
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                text_parts = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text().strip()
                        if text:
                            text_parts.append(text)
                    except Exception:
                        continue
                
                return "\n\n".join(text_parts) if text_parts else "PDF пустой"
                
            elif file_name.endswith('.rtf'):
                uploaded_file.seek(0)
                rtf_content = uploaded_file.read().decode('utf-8', errors='ignore')
                try:
                    text = rtf_to_text(rtf_content)
                    return text if text.strip() else "RTF файл пустой"
                except Exception as e:
                    return f"Ошибка парсинга RTF: {e}"
                    
            elif file_name.endswith('.txt'):
                uploaded_file.seek(0)
                text = uploaded_file.read().decode('utf-8', errors='ignore')
                return text if text.strip() else "TXT файл пустой"
                
            else:
                uploaded_file.seek(0)
                text = uploaded_file.read().decode('utf-8', errors='ignore')
                return text if text.strip() else "Неподдерживаемый формат"
                
        except Exception as e:
            return f"ОШИБКА: {str(e)}"
    
    
    def create_evaluation_prompt(self, resume_text, vacancy_text, industry):
        """Улучшенный промпт с заранее определенной отраслью"""
        
        industry_context = {
            'tech': {
                'critical_factors': ['актуальный технический стек', 'опыт с современными архитектурами', 'практические навыки разработки', 'знание алгоритмов и структур данных'],
                'red_flags': ['устаревшие технологии без развития', 'отсутствие практического опыта', 'слабые алгоритмические навыки', 'нет понимания современных подходов'],
                'bonus_points': ['contribution в open source', 'лидерство в техкоманде', 'опыт архитектурных решений', 'менторство junior разработчиков', 'участие в tech-комьюнити']
            },
            'fintech': {
                'critical_factors': ['опыт в финансовых услугах', 'знание регуляторных требований', 'понимание финансовых рисков', 'опыт работы с платежными системами'],
                'red_flags': ['отсутствие финансового опыта', 'незнание compliance требований', 'нет понимания финансовых рисков', 'только теоретические знания'],
                'bonus_points': ['сертификаты FRM/CFA/PRM', 'опыт в топ банках', 'знание международных стандартов', 'опыт внедрения финтех решений']
            },
            'healthcare': {
                'critical_factors': ['медицинское образование/лицензии', 'опыт работы с пациентами', 'знание медицинских протоколов', 'понимание этики здравоохранения'],
                'red_flags': ['отсутствие медицинского образования', 'нет опыта с пациентами', 'незнание медицинских стандартов', 'нарушения этики'],
                'bonus_points': ['специализированные сертификаты', 'научные публикации', 'опыт клинических исследований', 'международный медицинский опыт']
            },
            'retail': {
                'critical_factors': ['опыт работы с клиентами', 'понимание потребительского поведения', 'навыки продаж', 'знание товарных категорий'],
                'red_flags': ['отсутствие клиентского опыта', 'нет понимания розничных процессов', 'слабые коммуникативные навыки', 'незнание digital retail'],
                'bonus_points': ['опыт запуска новых продуктов', 'знание omnichannel', 'аналитические навыки', 'опыт управления категориями']
            },
            'manufacturing': {
                'critical_factors': ['опыт производственных процессов', 'знание качества и безопасности', 'понимание lean/six sigma', 'техническое образование'],
                'red_flags': ['отсутствие производственного опыта', 'незнание стандартов качества', 'нет понимания промышленной безопасности', 'только теоретические знания'],
                'bonus_points': ['сертификаты lean/six sigma', 'опыт автоматизации', 'знание IoT/Industry 4.0', 'опыт международного производства']
            },
            'consulting': {
                'critical_factors': ['аналитические навыки', 'опыт работы с клиентами', 'стратегическое мышление', 'презентационные навыки'],
                'red_flags': ['слабые аналитические способности', 'отсутствие клиентского опыта', 'нет стратегического мышления', 'плохие коммуникативные навыки'],
                'bonus_points': ['MBA топ школ', 'опыт в топ консалтинге', 'отраслевая экспертиза', 'международный опыт', 'опыт цифровой трансформации']
            },
            'education': {
                'critical_factors': ['педагогическое образование', 'опыт преподавания', 'знание методик обучения', 'способность к адаптации программ'],
                'red_flags': ['отсутствие педагогического опыта', 'незнание современных методик', 'слабые презентационные навыки', 'нет понимания цифрового обучения'],
                'bonus_points': ['ученая степень', 'международные сертификаты', 'опыт e-learning', 'научные публикации', 'инновационные методики']
            },
            'government': {
                'critical_factors': ['знание госслужбы', 'понимание регламентов', 'опыт работы с документооборотом', 'знание законодательства'],
                'red_flags': ['отсутствие госопыта', 'незнание нормативной базы', 'нет понимания бюрократических процессов', 'проблемы с документооборотом'],
                'bonus_points': ['юридическое образование', 'опыт в федеральных органах', 'знание международного права', 'опыт цифровизации госуслуг']
            }
        }
        
        context = industry_context.get(industry, {
            'critical_factors': ['релевантный опыт', 'соответствие уровня позиции', 'культурное соответствие'],
            'red_flags': ['частая смена работы', 'неясная мотивация', 'завышенные ожидания'],
            'bonus_points': ['стабильная карьера', 'дополнительное образование', 'лидерские качества']
        })
        
        return f"""
Ты - ведущий HR-эксперт с 20+ годами опыта в подборе персонала в отрасли {industry.upper()}.

=== КОНТЕКСТ ОТРАСЛИ: {industry.upper()} ===
КРИТИЧЕСКИЕ ФАКТОРЫ: {', '.join(context['critical_factors'])}
КРАСНЫЕ ФЛАГИ: {', '.join(context['red_flags'])}
БОНУСНЫЕ БАЛЛЫ: {', '.join(context['bonus_points'])}

=== ВАКАНСИЯ ===
{vacancy_text}

=== РЕЗЮМЕ КАНДИДАТА ===  
{resume_text}

=== ЭКСПЕРТНЫЙ АНАЛИЗ ===

Проведи профессиональную оценку по расширенной методологии:

**ЭТАП 1: ИЗВЛЕЧЕНИЕ КОНТАКТНОЙ ИНФОРМАЦИИ**
Внимательно проанализируй резюме и найди:
- Email адрес кандидата
- Номер телефона (в любом формате)
- Telegram (username, ссылки t.me, telegram.me, или @username)

**ЭТАП 2: ГЛУБОКИЙ АНАЛИЗ ВАКАНСИИ**
- Ключевые must-have требования (нарушение = автоотказ)
- Nice-to-have требования (влияют на общий балл)
- Скрытые требования (читай между строк)
- Уровень сложности задач (junior/middle/senior/expert)
- Командная динамика (лидерство/исполнение)

**ЭТАП 3: ПРОФИЛИРОВАНИЕ КАНДИДАТА**  
- Фактические достижения vs заявленные навыки
- Глубина экспертизы (поверхностное/среднее/экспертное)
- Траектория развития (рост/стагнация/деградация)
- Мотивационные факторы
- Культурное соответствие команде

**ЭТАП 4: МНОГОУРОВНЕВОЕ СОПОСТАВЛЕНИЕ**

**ОБЯЗАТЕЛЬНЫЕ ТРЕБОВАНИЯ (ВЕС: 40%)**
- Критические навыки (вето-фактор)
- Минимальный опыт 
- Обязательное образование/сертификация
[Если не выполнены - максимум 30 баллов общий]

**ПРОФЕССИОНАЛЬНАЯ ЭКСПЕРТИЗА (ВЕС: 35%)**  
- Глубина технических знаний
- Качество реализованных проектов
- Способность решать сложные задачи
- Инновационный подход

**ОТРАСЛЕВАЯ СПЕЦИАЛИЗАЦИЯ (ВЕС: 20%)**
- Знание специфики {industry}
- Понимание бизнес-процессов отрасли  
- Опыт решения отраслевых задач
- Соответствие регуляторным требованиям

**АДАПТИВНОСТЬ И РАЗВИТИЕ (ВЕС: 5%)**
- Способность к обучению
- Адаптация к изменениям
- Лидерский потенциал

**ЭТАП 5: RISK ASSESSMENT**
- Вероятность успешного прохождения испытательного срока
- Риски конфликтов в команде  
- Вероятность долгосрочного удержания
- Скорость выхода на полную эффективность

**СПЕЦИАЛЬНЫЕ ПРАВИЛА ОЦЕНКИ:**

🔴 **КРИТИЧЕСКИЕ НАРУШЕНИЯ (АВТОМАТИЧЕСКИЙ ОТКАЗ):**
- Отсутствие обязательных навыков из must-have
- Неправдивая информация в резюме
- Критическое несоответствие уровня позиции

🟡 **СЕРЬЕЗНЫЕ НЕДОСТАТКИ (СИЛЬНОЕ СНИЖЕНИЕ БАЛЛА):**
- Слабый отраслевой опыт при высоких требованиях
- Устаревшие навыки для tech-позиций
- Частая смена работы без логичного объяснения

🟢 **ПРЕИМУЩЕСТВА (ПОВЫШЕНИЕ БАЛЛА):**
- Превышение требований по ключевым навыкам
- Уникальная экспертиза
- Лидерский опыт и наставничество

ВЕРНИ СТРОГО JSON:
{{
    "contact_information": {{
        "email": "найденный email или null",
        "phone": "найденный телефон или null", 
        "telegram": "найденный telegram или null"
    }},
    "industry_analysis": {{
        "detected_industry": "{industry}",
        "industry_specific_requirements": ["требование 1", "требование 2"],
        "critical_success_factors": ["фактор 1", "фактор 2"],
        "industry_red_flags_found": ["флаг 1", "флаг 2"] 
    }},
    "vacancy_deep_dive": {{
        "must_have_requirements": ["обязательное 1", "обязательное 2"],
        "nice_to_have_requirements": ["желательное 1", "желательное 2"],
        "hidden_requirements": ["скрытое 1", "скрытое 2"],
        "position_complexity": "junior/middle/senior/expert",
        "team_dynamics_expected": "leadership/collaboration/independent"
    }},
    "candidate_profile": {{
        "total_experience_years": "конкретное число лет",
        "relevant_experience_years": "конкретное число лет",
        "expertise_level": "поверхностный/средний/экспертный",
        "career_trajectory": "растущий/стабильный/снижающийся",
        "key_achievements": ["достижение 1", "достижение 2"],
        "potential_concerns": ["беспокойство 1", "беспокойство 2"]
    }},
    "detailed_scoring": {{
        "mandatory_requirements": {{
            "score": 0-100,
            "weight": 40,
            "critical_violations": ["нарушение 1"],
            "met_requirements": ["выполненное 1", "выполненное 2"],
            "reasoning": "детальное обоснование по обязательным требованиям"
        }},
        "professional_expertise": {{
            "score": 0-100,
            "weight": 35,
            "technical_depth": "поверхностный/средний/экспертный",
            "project_quality": "низкое/среднее/высокое",
            "reasoning": "оценка профессиональной экспертизы с примерами"
        }},
        "industry_specialization": {{
            "score": 0-100,
            "weight": 20,
            "domain_match": "слабое/частичное/точное",
            "transferability_risk": "высокий/средний/низкий",
            "reasoning": "анализ отраслевой специализации"
        }},
        "adaptability": {{
            "score": 0-100,
            "weight": 5,
            "learning_ability": "низкая/средняя/высокая",
            "change_management": "слабое/среднее/сильное",
            "reasoning": "оценка адаптивности и потенциала роста"
        }}
    }},
    "risk_assessment": {{
        "probation_success_probability": "0-100%",
        "team_integration_risk": "низкий/средний/высокий",
        "retention_probability": "0-100%",
        "time_to_productivity": "1-3 месяца / 3-6 месяцев / 6+ месяцев",
        "overall_hiring_risk": "низкий/средний/высокий"
    }},
    "final_evaluation": {{
        "total_score": 0-100,
        "weighted_calculation": "40*(mandatory) + 35*(expertise) + 20*(industry) + 5*(adaptability) = итог",
        "recommendation": "strong_hire/hire/conditional_hire/weak_hire/no_hire",
        "confidence_level": "very_high/high/medium/low/very_low",
        "key_strengths": ["сильнейшая сторона 1", "сильнейшая сторона 2"],
        "critical_concerns": ["критическая проблема 1", "критическая проблема 2"],
        "hiring_rationale": "развернутое обоснование решения с учетом отраслевой специфики",
        "onboarding_strategy": "рекомендации по адаптации кандидата",
        "development_plan": "план развития на первый год",
        "compensation_analysis": "анализ соответствия ожиданий рынку"
    }}
}}
"""

    def evaluate_candidate(self, resume_text, vacancy_text, industry):
        """Отправка запроса к GPT-5 с заранее определенной отраслью"""
        try:
            response = self.client.chat.completions.create(
                model="gpt-5-mini",
                messages=[
                    {"role": "system", "content": "Ты - элитный HR-эксперт международного уровня с глубокой экспертизой в оценке талантов."},
                    {"role": "user", "content": self.create_evaluation_prompt(resume_text, vacancy_text, industry)}
                ],
            )
            result_text = response.choices[0].message.content
            
            start = result_text.find('{')
            end = result_text.rfind('}') + 1
            if start != -1 and end != 0:
                json_str = result_text[start:end]
                return json.loads(json_str)
            else:
                return {
                    "error": "Не удалось извлечь JSON из ответа",
                    "raw_response": result_text
                }
                
        except Exception as e:
            return {"error": f"Ошибка API: {str(e)}"}

    def create_final_evaluation_prompt(self, vacancy_text, resume_text, interview_log):
        """НОВЫЙ ПРОМПТ для финального заключения"""
        return f"""
Ты - главный HR-директор с 20-летним опытом, принимающий финальное решение о найме.
Тебе предоставлены все данные о кандидате: описание вакансии, его резюме, первичный AI-анализ и полный протокол технического интервью, проведенного AI-ассистентом.

ТВОЯ ЗАДАЧА: Проанализировать ВСЕ предоставленные материалы и дать финальное, комплексное заключение.

=== ВАКАНСИЯ ===
{vacancy_text}

=== РЕЗЮМЕ КАНДИДАТА ===
{resume_text}

=== ПРОТОКОЛ ТЕХНИЧЕСКОГО ИНТЕРВЬЮ ===
{interview_log}

=== АНАЛИЗ ===
Проведи глубокий анализ, ответив на следующие вопросы:
1.  **Соответствие резюме и интервью:** Подтвердились ли сильные стороны, заявленные в резюме, в ходе интервью? Были ли выявлены расхождения?
2.  **Техническая глубина:** Насколько глубоко кандидат понимает технологии, о которых говорит? Приводит ли он конкретные примеры или отвечает поверхностно?
3.  **Soft Skills:** Как кандидат ведет себя в диалоге? Уверенно? Структурированно? Способен ли признавать незнание? Как реагирует на сложные вопросы?
4.  **Ключевые "зеленые флаги":** Какие моменты в интервью однозначно говорят в пользу кандидата (например, успешное решение сложной задачи, демонстрация глубокой экспертизы)?
5.  **Ключевые "красные флаги":** Какие ответы или моменты вызывают наибольшее беспокойство?
6.  **Итоговая рекомендация:** Учитывая все "за" и "против", какое твое финальное решение?

=== ФОРМАТ ОТВЕТА ===
Верни ТОЛЬКО JSON объект следующей структуры:
{{
    "evaluation_summary": "Твое развернутое заключение в нескольких абзацах, где ты отвечаешь на все 6 вопросов анализа. Говори профессионально, но понятно.",
    "final_recommendation": "ОДНО ИЗ: 'hire' (нанять) или 'no_hire' (отказать)"
}}
"""

    def get_final_evaluation(self, vacancy_text, resume_text, interview_log):
        """Получить финальное заключение от GPT"""
        try:
            prompt = self.create_final_evaluation_prompt(vacancy_text, resume_text, interview_log)
            response = self.client.chat.completions.create(
                model="gpt-5-mini",
                messages=[{"role": "user", "content": prompt}],
                response_format={"type": "json_object"}
            )
            result_text = response.choices[0].message.content
            return json.loads(result_text)
        except Exception as e:
            return {"error": f"Ошибка API при получении финального заключения: {e}"}

def create_score_gauge(score, title):
    """Создание круговой диаграммы оценки"""
    fig = go.Figure(go.Indicator(
        mode = "gauge+number",
        value = score,
        domain = {'x': [0, 1], 'y': [0, 1]},
        title = {'text': title, 'font': {'size': 16}},
        gauge = {
            'axis': {'range': [None, 100]},
            'bar': {'color': "darkblue"},
            'steps': [
                {'range': [0, 50], 'color': "lightcoral"},
                {'range': [50, 75], 'color': "gold"}, 
                {'range': [75, 100], 'color': "lightgreen"}
            ],
            'threshold': {
                'line': {'color': "red", 'width': 4},
                'thickness': 0.75,
                'value': 80
            }
        }
    ))
    fig.update_layout(height=250, width=400, margin=dict(l=20, r=20, t=40, b=20))
    return fig

def display_contact_section(contacts, candidate_id, analysis_data, db_manager):
    """Отображение секции контактов с возможностью редактирования"""
    st.subheader("📞 Контактная информация")
    
    final_eval = analysis_data.get("final_evaluation", {})
    score = final_eval.get("total_score", 0)
    
    if score >= 50:
        if f"edit_email_{candidate_id}" not in st.session_state:
            st.session_state[f"edit_email_{candidate_id}"] = False
        if f"edit_phone_{candidate_id}" not in st.session_state:
            st.session_state[f"edit_phone_{candidate_id}"] = False
        if f"edit_telegram_{candidate_id}" not in st.session_state:
            st.session_state[f"edit_telegram_{candidate_id}"] = False
        
        col1, col2, col3 = st.columns(3)
        
        contact_info = analysis_data.get("contact_information", {})
        current_email = contacts.get('email') or contact_info.get('email')
        current_phone = contacts.get('phone') or contact_info.get('phone')  
        current_telegram = contacts.get('telegram') or contact_info.get('telegram')
        
        with col1:
            st.write("**Email:**")
            if current_email and not st.session_state[f"edit_email_{candidate_id}"]:
                col_email1, col_email2 = st.columns([4, 1])
                with col_email1:
                    st.success(f"✅ {current_email}")
                with col_email2:
                    if st.button("✏️", key=f"edit_email_btn_{candidate_id}", help="Редактировать"):
                        st.session_state[f"edit_email_{candidate_id}"] = True
                        st.rerun()
                new_email = current_email
            else:
                new_email = st.text_input("Email:", value=current_email or "", key=f"email_input_{candidate_id}", placeholder="example@mail.com")
        
        with col2:
            st.write("**Телефон:**")
            if current_phone and not st.session_state[f"edit_phone_{candidate_id}"]:
                col_phone1, col_phone2 = st.columns([4, 1])
                with col_phone1:
                    st.success(f"✅ {current_phone}")
                with col_phone2:
                    if st.button("✏️", key=f"edit_phone_btn_{candidate_id}", help="Редактировать"):
                        st.session_state[f"edit_phone_{candidate_id}"] = True
                        st.rerun()
                new_phone = current_phone
            else:
                new_phone = st.text_input("Телефон:", value=current_phone or "", key=f"phone_input_{candidate_id}", placeholder="+7 999 123-45-67")
        
        with col3:
            st.write("**Telegram:**")
            if current_telegram and not st.session_state[f"edit_telegram_{candidate_id}"]:
                col_tg1, col_tg2 = st.columns([4, 1])
                with col_tg1:
                    st.success(f"✅ {current_telegram}")
                with col_tg2:
                    if st.button("✏️", key=f"edit_telegram_btn_{candidate_id}", help="Редактировать"):
                        st.session_state[f"edit_telegram_{candidate_id}"] = True
                        st.rerun()
                new_telegram = current_telegram
            else:
                new_telegram = st.text_input("Telegram:", value=current_telegram or "", key=f"telegram_input_{candidate_id}", placeholder="@username")
        
        if st.button("💾 Сохранить контакты", key=f"save_contacts_{candidate_id}"):
            db_manager.update_candidate_contacts(candidate_id, new_email, new_phone, new_telegram)
            
            st.session_state[f"edit_email_{candidate_id}"] = False
            st.session_state[f"edit_phone_{candidate_id}"] = False
            st.session_state[f"edit_telegram_{candidate_id}"] = False
            
            st.success("Контакты обновлены!")
            st.rerun()
        
        return {
            'email': new_email,
            'phone': new_phone, 
            'telegram': new_telegram
        }
    
    return None

def display_interview_toggle(analysis_id, current_status, contacts, score, db_manager, analysis):
    """Отображение переключателя статуса собеседования с отправкой запроса на сервер."""
    st.subheader("🎯 Статус собеседования")
    
    if score >= 50:
        
        has_contacts = (contacts and (contacts.get('email') or contacts.get('telegram')))
        
        if not has_contacts:
            st.warning("⚠️ Для назначения собеседования необходимо заполнить Email или Telegram кандидата")
            return

        col1, col2 = st.columns([3, 1])
        
        with col1:
            current_status_text = "Назначено" if current_status else "Не назначено"
            status_color = "🟢" if current_status else "🔴"
            st.write(f"{status_color} **Собеседование:** {current_status_text}")
        
        with col2:
            new_status = st.toggle(
                "Назначить собеседование",
                value=current_status,
                key=f"interview_toggle_{analysis_id}"
            )
        
        if new_status != current_status:
            db_manager.update_interview_status(analysis_id, new_status)
            
            if new_status:
                with st.spinner("Отправка команды на основной сервер..."):
                    try:
                        payload = {
                            "webhook_type": "interview_approved",
                            "analysis_id": analysis_id,
                            "contacts": contacts,
                            "vacancy": {
                                "title": analysis['vacancy_title'],
                                "content": analysis['vacancy_content']
                            }
                        }
                        
                        response = requests.post(MAIN_SERVER_URL, json=payload, timeout=10)
                        
                        if response.status_code == 200:
                            st.success("✅ Команда на запуск интервью успешно отправлена!")
                        else:
                            st.error(f"❌ Ошибка сервера: {response.status_code} - {response.text}")
                    
                    except requests.exceptions.RequestException as e:
                        st.error(f"❌ Не удалось подключиться к основному серверу: {e}")
            
            st.rerun()

def display_results(evaluation, show_full=True):
    """Отображение результатов анализа (только JSON часть)"""
    if "error" in evaluation:
        st.error(f"Ошибка: {evaluation['error']}")
        if "raw_response" in evaluation:
            with st.expander("Посмотреть ответ нейросети"):
                st.text(evaluation["raw_response"])
        return None
    
    russian_names = {
                    'mandatory_requirements': 'Обязательные требования',
                    'professional_expertise': 'Профессиональная экспертиза', 
                    'industry_specialization': 'Отраслевая специализация',
                    'adaptability': 'Адаптивность'
                    }
    
    final_eval = evaluation.get("final_evaluation", {})
    detailed_scoring = evaluation.get("detailed_scoring", {})
    
    score = final_eval.get("total_score", 0)
    
    if score >= 75:
        recommendation = "strong_hire"
    elif score >= 50:
        recommendation = "hire"
    else:
        recommendation = "no_hire"
    
    confidence = final_eval.get("confidence_level", "medium")
    
    if show_full:
        col1, col2, col3 = st.columns([2, 2, 3])
        
        with col1:
            st.plotly_chart(create_score_gauge(score, f"Оценка резюме: {score}/100"))
        
        with col2:
            rec_map = {
                "strong_hire": ("✅ НАСТОЯТЕЛЬНО РЕКОМЕНДУЮ", "НАНЯТЬ НЕМЕДЛЕННО"),
                "hire": ("✅ РЕКОМЕНДУЮ", "НАНЯТЬ"),
                "no_hire": ("❌ НЕ РЕКОМЕНДУЮ", "ОТКАЗАТЬ")
            }
            
            rec_text, rec_action = rec_map.get(recommendation, ("❓ НЕОПРЕДЕЛЕНО", "ТРЕБУЕТ АНАЛИЗ"))
            
            if recommendation in ["strong_hire", "hire"]:
                st.success(rec_text)
                st.success(f"**{rec_action}**")
            else:
                st.error(rec_text)
                st.error(f"**{rec_action}**")
                
            confidence_map = {
                "very_high": "очень высокая", "high": "высокая", "medium": "средняя", 
                "low": "низкая", "very_low": "очень низкая"
            }
            confidence_emoji = {
                "very_high": "🎯", "high": "🎯", "medium": "🤔", "low": "❓", "very_low": "❓"
            }
            conf_ru = confidence_map.get(confidence, "средняя")
            st.info(f"{confidence_emoji.get(confidence, '🤔')} Уверенность: {conf_ru}")
        
        with col3:
            st.subheader("💭 Логика решения")
            decision_logic = final_eval.get("hiring_rationale", "Обоснование недоступно")
            st.write(decision_logic)
        
        st.markdown("---")
        
        if "industry_analysis" in evaluation:
            industry_info = evaluation["industry_analysis"]
            st.subheader(f"🏢 Анализ отрасли: {industry_info.get('detected_industry', 'general').title()}")
            col1, col2 = st.columns(2)
            with col1:
                st.write("**Специфические требования отрасли:**")
                for req in industry_info.get("industry_specific_requirements", []):
                    st.write(f"• {req}")
            with col2:
                if industry_info.get("industry_red_flags_found"):
                    st.write("**Обнаруженные красные флаги:**")
                    for flag in industry_info["industry_red_flags_found"]:
                        st.write(f"🚩 {flag}")
            st.markdown("---")
        
        st.subheader("📊 Детальный анализ по критериям")
        if detailed_scoring:
            scores_data = []
            for criterion, data in detailed_scoring.items():
                if isinstance(data, dict) and "score" in data:
                    weight = data.get('weight', 0)
                    
                    criterion_name = russian_names.get(criterion, criterion.replace('_', ' ').title())
                    scores_data.append({
                        "Критерий": f"{criterion_name}\n({weight}%)",
                        "Оценка": data["score"],
                        "Вес": weight
                    })
            if scores_data:
                fig = px.bar(scores_data, x="Критерий", y="Оценка", title="Оценки по критериям",
                             color="Оценка", color_continuous_scale="RdYlGn", range_color=[0, 100])
                fig.update_layout(height=400)
                st.plotly_chart(fig, use_container_width=True)
        
        if "risk_assessment" in evaluation:
            risk_info = evaluation["risk_assessment"]
            st.subheader("⚠️ Оценка рисков найма")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Успех испытательного срока", risk_info.get("probation_success_probability", "N/A"))
                st.metric("Удержание сотрудника", risk_info.get("retention_probability", "N/A"))
            with col2:
                st.metric("Риск интеграции в команду", risk_info.get("team_integration_risk", "средний"))
                st.metric("Время выхода на эффективность", risk_info.get("time_to_productivity", "3-6 месяцев"))
            with col3:
                overall_risk = risk_info.get("overall_hiring_risk", "средний")
                if overall_risk == "низкий": st.success(f"Общий риск найма: {overall_risk}")
                elif overall_risk == "средний": st.warning(f"Общий риск найма: {overall_risk}")
                else: st.error(f"Общий риск найма: {overall_risk}")
        
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("💪 Ключевые сильные стороны")
            strengths = final_eval.get("key_strengths", [])
            for strength in strengths: st.write(f"• {strength}")
            if not strengths: st.write("Не выявлены")
        with col2:
            st.subheader("⚠️ Критические замечания")
            concerns = final_eval.get("critical_concerns", [])
            for concern in concerns: st.write(f"• {concern}")
            if not concerns: st.write("Отсутствуют")
        
        if "development_plan" in final_eval:
            st.subheader("📈 План развития кандидата")
            st.info(final_eval["development_plan"])
        
        st.subheader("🔍 Подробный анализ по критериям")
        for criterion, data in detailed_scoring.items():
            if isinstance(data, dict):
                score = data.get('score', 0)
                reasoning = data.get('reasoning', 'Анализ недоступен')
                if score >= 80: color = "🟢"
                elif score >= 60: color = "🟡"
                else: color = "🔴"
                criterion_name = russian_names.get(criterion, criterion.replace('_', ' ').title())
                with st.expander(f"{color} {criterion_name} - {score}/100"):
                    st.write(reasoning)
    
    return {
        "score": score,
        "recommendation": recommendation,
        "confidence": confidence
    }

def show_history_sidebar(db_manager):
    """Боковое меню с кнопкой возврата и индикаторами интервью"""
    
    if st.sidebar.button("🏠 Вернуться к анализу", use_container_width=True):
        st.session_state.current_view = 'main'
        st.session_state.selected_analysis_id = None
        st.rerun()
    
    st.sidebar.markdown("---")
    st.sidebar.header("📚 История анализов")
    
    search_term = st.sidebar.text_input("🔍 Поиск анализов", placeholder="Название вакансии или кандидат...")
    
    if search_term:
        analyses = db_manager.search_analyses(search_term)
    else:
        analyses = db_manager.get_all_analyses()
    
    if not analyses:
        st.sidebar.info("Анализы не найдены")
        return None
    
    total_analyses = len(analyses)
    
    hired_count = sum(1 for a in analyses if (a['total_score'] or 0) >= 50)
    
    avg_score = sum(a['total_score'] or 0 for a in analyses) / total_analyses if total_analyses > 0 else 0
    
    st.sidebar.metric("Всего анализов", total_analyses)
    st.sidebar.metric("Рекомендованы к найму", hired_count)  
    st.sidebar.metric("Средний балл", f"{avg_score:.1f}")
    
    st.sidebar.markdown("---")
    
    selected_analysis = None
    
    for analysis in analyses[:20]:
        created_at = datetime.fromisoformat(analysis['created_at'].replace('Z', '+00:00')).strftime("%d.%m.%Y %H:%M")
        
        if (analysis['total_score'] or 0) >= 50:
            emoji = "✅"
        else:
            emoji = "❌"
        
        interview_emoji = "🎙️" if analysis['interview_id'] else ""
        
        final_eval_emoji = "🏆" if 'final_evaluation_id' in analysis.keys() and analysis['final_evaluation_id'] else ""
        
        button_text = f"{emoji}{interview_emoji}{final_eval_emoji} {analysis['candidate_name'][:15]}...\n{analysis['vacancy_title'][:20]}...\n{created_at}"
        
        if st.sidebar.button(button_text, key=f"analysis_{analysis['id']}", help=f"Балл: {analysis['total_score']}/100"):
            selected_analysis = analysis['id']
    
    return selected_analysis

def display_interview_report(analysis_data, interview_data, qa_data, db_manager, evaluator):
    """Отображение полного отчета по собеседованию"""
    try:
        final_scores = json.loads(interview_data.get('final_scores', '{}'))
        phase_breakdown = json.loads(interview_data.get('phase_breakdown', '{}'))
        
        overall_score = final_scores.get('overall_score', 0)
        
        resume_score = analysis_data['total_score'] if 'total_score' in analysis_data.keys() else 0

        if overall_score >= 50 and resume_score >= 50:
            recommendation_text = "hire"
        else:
            recommendation_text = "no_hire"
        
        col1, col2, col3 = st.columns([2, 1, 2])
        with col1:
            st.plotly_chart(create_score_gauge(overall_score, "Оценка за интервью"))
        with col2:
            st.metric("Длительность", f"{interview_data.get('duration_seconds', 0) // 60} мин")
            st.metric("Задано вопросов", interview_data.get('total_questions', 0))
        with col3:
            rec_map = {
                "hire": ("✅ КАНДИДАТ ПОДХОДИТ", "success"),
                "no_hire": ("❌ КАНДИДАТ НЕ ПОДХОДИТ", "error")
            }
            rec_text, rec_type = rec_map.get(recommendation_text, ("❓ Неопределено", "info"))
            
            if rec_type == "success":
                st.success(rec_text)
                st.info(f"Оценка за резюме: {resume_score}/100\nОценка за интервью: {overall_score}/100")
            else:
                st.error(rec_text)
                st.warning(f"Оценка за резюме: {resume_score}/100\nОценка за интервью: {overall_score}/100")

        st.markdown("---")

        st.subheader("🏆 Финальное заключение по кандидату")
        
        final_evaluation = db_manager.get_final_evaluation(analysis_data['id'])

        if final_evaluation:
            st.success("Заключение было сгенерировано ранее:")
            st.markdown(final_evaluation['evaluation_summary'])
            if final_evaluation['final_recommendation'] == 'hire':
                st.success("Итоговая рекомендация: **Нанять**")
            else:
                st.error("Итоговая рекомендация: **Отказать**")
        
        if st.button("🤖 Сгенерировать/обновить финальное заключение", key="generate_final_eval"):
            with st.spinner("Анализирую все данные... Это может занять минуту..."):
                vacancy_text = analysis_data['vacancy_content']
                resume_text = analysis_data['candidate_resume']
                
                interview_log_parts = []
                for qa in qa_data:
                    interview_log_parts.append(f"Вопрос {qa['question_number']}: {qa['question_text']}")
                    interview_log_parts.append(f"Ответ: {qa['answer_text']}")
                    analysis_notes = qa['analysis_notes'] if 'analysis_notes' in qa.keys() and qa['analysis_notes'] else 'N/A'
                    interview_log_parts.append(f"Анализ AI: {analysis_notes}\n")
                interview_log = "\n".join(interview_log_parts)

                result = evaluator.get_final_evaluation(vacancy_text, resume_text, interview_log)
                
                if "error" not in result:
                    summary = result.get("evaluation_summary", "Не удалось сгенерировать заключение.")
                    recommendation = result.get("final_recommendation", "no_hire")
                    db_manager.save_final_evaluation(analysis_data['id'], summary, recommendation)
                    st.success("Финальное заключение сохранено!")
                    st.rerun()
                else:
                    st.error(result["error"])

        st.markdown("---")

        st.subheader("📈 Анализ по фазам интервью")
        if phase_breakdown:
            phase_data = []
            phase_names = {
                'exploration': 'Исследование',
                'validation': 'Проверка', 
                'stress_test': 'Стресс-тест',
                'soft_skills': 'Софт скилы',
                'wrap_up': 'Завершение'
            }
            for phase, stats in phase_breakdown.items():
                phase_data.append({
                "Фаза": phase_names.get(phase, phase.replace('_', ' ').title()),
                "Средний балл": stats.get('avg_score', 0),
                "Вопросов": stats.get('questions_asked', 0)
            })
            
            df = pd.DataFrame(phase_data)
            fig = px.bar(df, x="Фаза", y="Средний балл", text="Вопросов",
                         title="Эффективность по фазам интервью",
                         color="Средний балл", color_continuous_scale="RdYlGn",
                         range_color=[0, 10])
            fig.update_traces(texttemplate='%{text} вопр.', textposition='outside')
            st.plotly_chart(fig, use_container_width=True)

        st.markdown("---")

        with st.expander("📝 Показать полный протокол собеседования (Q&A)"):
            if not qa_data:
                st.info("Детальный протокол отсутствует.")
            for qa in qa_data:
                st.markdown(f"**Вопрос {qa['question_number']} (Фаза: {qa['question_phase']}, Сложность: {qa['question_difficulty']}):**")
                st.markdown(f"> {qa['question_text']}")
                st.markdown("**Ответ кандидата:**")
                st.text_area("", value=qa['answer_text'], height=100, disabled=True, key=f"ans_{qa['id']}")
                
                st.markdown("**Анализ ответа:**")
                analysis_cols = st.columns(4)
                analysis_cols[0].metric("Тех. навыки", f"{qa['technical_score']}/10")
                analysis_cols[1].metric("Коммуникация", f"{qa['communication_score']}/10")
                analysis_cols[2].metric("Уверенность", f"{qa['confidence_score']}/10")
                analysis_cols[3].metric("Глубина", f"{qa['depth_score']}/10")
                
                analysis_notes = qa['analysis_notes'] if 'analysis_notes' in qa.keys() and qa['analysis_notes'] else None
                if analysis_notes:
                    st.info(f"**Заметки интервьюера:** {analysis_notes}")
                st.markdown("---")

    except (json.JSONDecodeError, KeyError) as e:
        st.error(f"Не удалось отобразить отчет по собеседованию. Ошибка данных: {e}")

def show_analysis_details(db_manager, analysis_id):
    """Показать детали конкретного анализа с вкладкой для интервью"""
    analysis = db_manager.get_analysis_by_id(analysis_id)
    evaluator = CandidateEvaluator()
    
    if not analysis:
        st.error("Анализ не найден")
        return
    
    st.header(f"📄 Анализ: {analysis['candidate_name']}")
    st.subheader(f"Вакансия: {analysis['vacancy_title']}")
    
    created_at = datetime.fromisoformat(analysis['created_at'].replace('Z', '+00:00')).strftime("%d.%m.%Y в %H:%M")
    st.caption(f"Создан: {created_at}")
    
    st.markdown("---")
    
    interview_data, qa_data = db_manager.get_interview_by_analysis_id(analysis_id)
    
    if interview_data:
        tab1, tab2 = st.tabs(["📄 Анализ резюме", "🎙️ Отчет по собеседованию"])
        
        with tab1:
            display_resume_analysis(analysis, db_manager)
            
        with tab2:
            display_interview_report(analysis, interview_data, qa_data, db_manager, evaluator)
            
    else:
        display_resume_analysis(analysis, db_manager)

def display_resume_analysis(analysis, db_manager):
    """Отображение части, касающейся только анализа резюме"""
    try:
        result_json = json.loads(analysis['result_json'])
        
        display_results(result_json, show_full=True)
        
        final_eval = result_json.get("final_evaluation", {})
        score = final_eval.get("total_score", 0)

        if score >= 50:
            st.markdown("---")
            
            contacts = {
                'email': analysis['candidate_email'],
                'phone': analysis['candidate_phone'],
                'telegram': analysis['candidate_telegram']
            }
            candidate_id = analysis['candidate_id']
            
            updated_contacts = display_contact_section(contacts, candidate_id, result_json, db_manager)
            
            final_contacts = updated_contacts or contacts
            current_interview_status = analysis['needs_interview']
            
            display_interview_toggle(analysis['id'], current_interview_status, final_contacts, score, db_manager, dict(analysis))
        
    except json.JSONDecodeError:
        st.error("Ошибка при загрузке данных анализа")
    
    with st.expander("📋 Исходные документы"):
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Описание вакансии")
            st.text_area("", value=analysis['vacancy_content'], height=300, disabled=True, key="vacancy_content")
        with col2:
            st.subheader("Резюме кандидата")  
            st.text_area("", value=analysis['candidate_resume'], height=300, disabled=True, key="candidate_resume")

def show_multiple_resumes_interface(db_manager, evaluator):
    """Интерфейс для множественных резюме"""
    st.header("👥 Анализ нескольких кандидатов для одной вакансии")
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    col1, col2 = st.columns([1, 2], gap="large")
    
    with col1:
        st.subheader("📋 Вакансия")
        st.markdown("<br>", unsafe_allow_html=True)
        
        vacancy_file = st.file_uploader(
            "Загрузите описание вакансии",
            type=['docx', 'pdf', 'txt', 'rtf'],
            key="multi_vacancy"
        )
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        vacancy_text = st.text_area(
            "Или введите описание:",
            height=200,
            key="multi_vacancy_text"
        )
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        vacancy_title = st.text_input("Название вакансии:", key="multi_vacancy_title")
    
    with col2:
        st.subheader("👤 Резюме кандидатов")
        st.markdown("<br>", unsafe_allow_html=True)
        
        resume_files = st.file_uploader(
            "Загрузите резюме (можно выбрать несколько файлов)",
            type=['docx', 'pdf', 'txt', 'rtf'],
            accept_multiple_files=True,
            key="multi_resumes"
        )
        
        if resume_files:
            st.success(f"Загружено резюме: {len(resume_files)} файлов")
            for i, file in enumerate(resume_files):
                st.write(f"{i+1}. {file.name}")
    
    st.markdown("<br><br>", unsafe_allow_html=True)
    
    if st.button("🚀 Проанализировать всех кандидатов", type="primary", use_container_width=True):
        if not (vacancy_file or vacancy_text.strip()):
            st.error("Загрузите описание вакансии")
            return
            
        if not resume_files:
            st.error("Загрузите резюме кандидатов")
            return
        
        if vacancy_file:
            vacancy_content = evaluator.extract_text_from_file(vacancy_file)
            v_title = vacancy_title or vacancy_file.name.split('.')[0]
        else:
            vacancy_content = vacancy_text.strip()
            v_title = vacancy_title or "Без названия"
        
        if vacancy_content.startswith("ОШИБКА"):
            st.error(vacancy_content)
            return
        
        with st.spinner("🔍 Определяю отрасль вакансии..."):
            industry = evaluator.classify_industry(vacancy_content)
            st.success(f"✅ Отрасль определена: {industry}")
        
        results = []
        vacancy_id = db_manager.save_vacancy(v_title, vacancy_content, industry)
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, resume_file in enumerate(resume_files):
            status_text.text(f"Анализируется: {resume_file.name} ({i+1}/{len(resume_files)})")
            
            resume_content = evaluator.extract_text_from_file(resume_file)
            
            if resume_content.startswith("ОШИБКА"):
                results.append({
                    'candidate_name': resume_file.name,
                    'error': resume_content,
                    'score': 0,
                    'recommendation': 'error'
                })
                continue
            
            evaluation = evaluator.evaluate_candidate(resume_content, vacancy_content, industry)
            
            if evaluation and "error" not in evaluation:
                final_eval = evaluation.get("final_evaluation", {})
                candidate_name = resume_file.name.split('.')[0]
                
                contact_info = evaluation.get("contact_information", {})
                
                score = final_eval.get("total_score", 0)
                needs_interview = score >= 50 and (contact_info.get('email') or contact_info.get('telegram'))
                
                candidate_id = db_manager.save_candidate(
                    candidate_name, 
                    resume_content,
                    contact_info.get('email'),
                    contact_info.get('phone'),
                    contact_info.get('telegram')
                )
                
                analysis_id = db_manager.save_analysis(
                    vacancy_id, 
                    candidate_id, 
                    evaluation,
                    score,
                    "hire" if score >= 50 else "no_hire",
                    final_eval.get("confidence_level", "medium"),
                    "batch_multiple_resumes",
                    needs_interview
                )
                
                results.append({
                    'candidate_name': candidate_name,
                    'analysis_id': analysis_id,
                    'score': score,
                    'recommendation': "hire" if score >= 50 else "no_hire",
                    'confidence': final_eval.get("confidence_level", "medium"),
                    'key_strengths': final_eval.get("key_strengths", []),
                    'critical_concerns': final_eval.get("critical_concerns", [])
                })
            else:
                results.append({
                    'candidate_name': resume_file.name,
                    'error': evaluation.get('error', 'Неизвестная ошибка'),
                    'score': 0,
                    'recommendation': 'error'
                })
            
            progress_bar.progress((i + 1) / len(resume_files))
        
        status_text.text("Анализ завершен!")
        
        st.subheader("📊 Сравнительная таблица кандидатов")
        
        results_sorted = sorted(results, key=lambda x: x.get('score', 0), reverse=True)
        
        table_data = []
        for i, result in enumerate(results_sorted):
            if 'error' in result:
                table_data.append({
                    'Место': f"{i+1}",
                    'Кандидат': result['candidate_name'],
                    'Балл': "❌ Ошибка",
                    'Рекомендация': "Ошибка анализа",
                    'Уверенность': "-"
                })
            else:
                rec_emoji = {
                    'hire': '✅ Рекомендую',
                    'no_hire': '❌ Не рекомендую'
                }
                
                conf_emoji = {
                    'very_high': '🎯', 'high': '🎯', 'medium': '🤔', 'low': '❓', 'very_low': '❓'
                }
                
                table_data.append({
                    'Место': f"🥇 {i+1}" if i == 0 else f"🥈 {i+1}" if i == 1 else f"🥉 {i+1}" if i == 2 else f"{i+1}",
                    'Кандидат': result['candidate_name'],
                    'Балл': f"{result['score']}/100",
                    'Рекомендация': rec_emoji.get(result['recommendation'], result['recommendation']),
                    'Уверенность': f"{conf_emoji.get(result['confidence'], '🤔')} {result['confidence']}"
                })
        
        df = pd.DataFrame(table_data)
        st.dataframe(df, use_container_width=True)
        
        st.subheader("🏆 Топ-3 кандидата")
        
        for i, result in enumerate(results_sorted[:3]):
            if 'error' not in result:
                with st.expander(f"{'🥇' if i==0 else '🥈' if i==1 else '🥉'} {result['candidate_name']} - {result['score']}/100"):
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.write("**Сильные стороны:**")
                        for strength in result.get('key_strengths', []):
                            st.write(f"• {strength}")
                    
                    with col2:
                        st.write("**Замечания:**")
                        for concern in result.get('critical_concerns', []):
                            st.write(f"• {concern}")
                    
                    if st.button(f"Подробный анализ {result['candidate_name']}", key=f"detail_{result['analysis_id']}"):
                        st.session_state.current_view = 'analysis_detail'
                        st.session_state.selected_analysis_id = result['analysis_id']
                        st.rerun()
        
        db_manager.save_batch_analysis(
            f"Batch анализ: {v_title}",
            "multiple_resumes",
            results
        )

def main():
    db_type = os.getenv('DATABASE_TYPE', 'sqlite')
    db_manager = DatabaseManager(db_type)
    evaluator = CandidateEvaluator()
    
    st.title("HR-Аватар: Автоматический анализ кандидатов")
    st.markdown("*Умный подбор персонала с помощью ИИ*")
    st.markdown("---")
    
    if st.session_state.current_view == 'analysis_detail':
        selected_analysis_id = show_history_sidebar(db_manager)
        
        if selected_analysis_id:
            st.session_state.selected_analysis_id = selected_analysis_id
            st.rerun()
        
        if st.session_state.selected_analysis_id:
            show_analysis_details(db_manager, st.session_state.selected_analysis_id)
        
        return
    
    selected_analysis_id = show_history_sidebar(db_manager)
    
    if selected_analysis_id:
        st.session_state.current_view = 'analysis_detail'
        st.session_state.selected_analysis_id = selected_analysis_id
        st.rerun()
    
    st.header("🆕 Новый анализ")
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    analysis_mode = st.selectbox(
        "**Выберите тип анализа:**",
        options=[
            "single",
            "multiple_resumes"
        ],
        format_func=lambda x: {
            "single": "📄 Одна вакансия → Одно резюме",
            "multiple_resumes": "📄 Одна вакансия → Несколько резюме"
        }[x],
        key="analysis_mode_select"
    )
    
    st.session_state.analysis_mode = analysis_mode
    
    st.markdown("<br><br>", unsafe_allow_html=True)
    
    if analysis_mode == "multiple_resumes":
        show_multiple_resumes_interface(db_manager, evaluator)
        
    else:
        st.header("📄 Стандартный анализ")
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        col1, col2 = st.columns(2, gap="large")
        
        with col1:
            st.subheader("📋 Описание вакансии")
            st.markdown("<br>", unsafe_allow_html=True)
            
            vacancy_file = st.file_uploader(
                "Загрузите файл с описанием вакансии",
                type=['docx', 'pdf', 'txt', 'rtf'],
                key="single_vacancy"
            )
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            vacancy_text = st.text_area(
                "Или введите описание вакансии вручную:",
                height=200,
                key="single_vacancy_text",
                placeholder="Вставьте текст вакансии сюда..."
            )
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            vacancy_title = st.text_input(
                "Название вакансии:",
                placeholder="Например: Senior Python Developer"
            )
        
        with col2:
            st.subheader("👤 Резюме кандидата")
            st.markdown("<br>", unsafe_allow_html=True)
            
            resume_file = st.file_uploader(
                "Загрузите резюме кандидата",
                type=['docx', 'pdf', 'txt', 'rtf'],
                key="single_resume"
            )
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            resume_text = st.text_area(
                "Или введите резюме вручную:",
                height=200,
                key="single_resume_text",
                placeholder="Вставьте текст резюме сюда..."
            )
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            candidate_name = st.text_input(
                "Имя кандидата:",
                placeholder="Например: Иван Петров"
            )
        
        st.markdown("<br><br>", unsafe_allow_html=True)
        
        if st.button("🚀 Проанализировать кандидата", type="primary", use_container_width=True):
            final_vacancy_text = ""
            final_resume_text = ""
            
            final_vacancy_title = vacancy_title
            final_candidate_name = candidate_name
            
            if vacancy_file:
                with st.spinner("Обрабатываю файл вакансии..."):
                    final_vacancy_text = evaluator.extract_text_from_file(vacancy_file)
                    
                    if final_vacancy_text.startswith("ОШИБКА"):
                        st.error(final_vacancy_text)
                        st.stop()
                    else:
                        st.success(f"✅ Вакансия обработана: {len(final_vacancy_text)} символов")
                        
                        if not final_vacancy_title:
                            final_vacancy_title = vacancy_file.name.split('.')[0]
                            
            elif vacancy_text.strip():
                final_vacancy_text = vacancy_text.strip()
                st.success("✅ Используется текст вакансии, введенный вручную")
                if not final_vacancy_title:
                    final_vacancy_title = "Вакансия без названия"
            
            if resume_file:
                with st.spinner("Обрабатываю резюме..."):
                    final_resume_text = evaluator.extract_text_from_file(resume_file)
                    
                    if final_resume_text.startswith("ОШИБКА"):
                        st.error(final_resume_text)
                        st.stop()
                    else:
                        st.success(f"✅ Резюме обработано: {len(final_resume_text)} символов")
                        
                        if not final_candidate_name:
                            final_candidate_name = resume_file.name.split('.')[0]
                            
            elif resume_text.strip():
                final_resume_text = resume_text.strip()
                st.success("✅ Используется текст резюме, введенный вручную")
                if not final_candidate_name:
                    final_candidate_name = "Кандидат без имени"
            
            if not final_vacancy_text:
                st.error("❌ Загрузите файл вакансии или введите описание вручную")
                st.stop()
                
            if not final_resume_text:
                st.error("❌ Загрузите резюме или введите текст вручную") 
                st.stop()
            
            with st.spinner("🔍 Определяю отрасль вакансии..."):
                industry = evaluator.classify_industry(final_vacancy_text)
                st.success(f"✅ Отрасль определена: {industry}")
            
            with st.spinner("🧠 Проводится глубокий анализ кандидата... Это займет пару минут."):
                evaluation = evaluator.evaluate_candidate(final_resume_text, final_vacancy_text, industry)
                
                if evaluation and "error" not in evaluation:
                    st.success("✅ Анализ завершен!")
                    st.markdown("---")
                    
                    try:
                        vacancy_id = db_manager.save_vacancy(
                            final_vacancy_title,
                            final_vacancy_text,
                            industry
                        )
                        
                        contact_info = evaluation.get("contact_information", {})
                        
                        candidate_id = db_manager.save_candidate(
                            final_candidate_name,
                            final_resume_text,
                            contact_info.get('email'),
                            contact_info.get('phone'),
                            contact_info.get('telegram')
                        )
                        
                        final_eval = evaluation.get("final_evaluation", {})
                        score = final_eval.get("total_score", 0)
                        recommendation = "hire" if score >= 50 else "no_hire"
                        needs_interview = score >= 50 and (contact_info.get('email') or contact_info.get('telegram'))
                        
                        analysis_id = db_manager.save_analysis(
                            vacancy_id,
                            candidate_id,
                            evaluation,
                            score,
                            recommendation,
                            final_eval.get("confidence_level", "medium"),
                            "single",
                            needs_interview
                        )
                        
                        contacts = {
                            'email': contact_info.get('email'),
                            'phone': contact_info.get('phone'),
                            'telegram': contact_info.get('telegram')
                        }
                        
                        display_results(evaluation)

                        final_eval = evaluation.get("final_evaluation", {})
                        score = final_eval.get("total_score", 0)

                        if score >= 50:
                            st.markdown("---")
                            
                            contacts = {
                                'email': contact_info.get('email'),
                                'phone': contact_info.get('phone'),
                                'telegram': contact_info.get('telegram')
                            }
                            
                            updated_contacts = display_contact_section(contacts, candidate_id, evaluation, db_manager)
                            
                            final_contacts = updated_contacts or contacts
                            
                            analysis_for_toggle = {
                                'id': analysis_id,
                                'vacancy_title': final_vacancy_title,
                                'vacancy_content': final_vacancy_text
                            }
                            display_interview_toggle(analysis_id, needs_interview, final_contacts, score, db_manager, analysis_for_toggle)
                        
                        st.info(f"💾 Анализ сохранен в базе данных (ID: {analysis_id[:8]}...)")
                        
                    except Exception as e:
                        st.warning(f"Анализ выполнен, но не удалось сохранить в БД: {e}")
                        display_results(evaluation)

if __name__ == "__main__":
    main()